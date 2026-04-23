import os
import json
import pdfplumber
from docx import Document
import requests

OLLAMA_URL = "http://localhost:11434/api/generate"
MODEL = "llama3" # Defaulting to llama3

def extract_text_from_pdf(pdf_path):
    if not os.path.exists(pdf_path):
        return None
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

def ollama_generate(prompt):
    payload = {
        "model": MODEL,
        "prompt": prompt,
        "stream": False
    }
    try:
        response = requests.post(OLLAMA_URL, json=payload, timeout=60)
        return response.json().get("response", "Error generating text.")
    except Exception as e:
        return f"Error connecting to Ollama: {e}"

def generate_tailored_documents(job_details):
    company = job_details.get("company", "Unknown Company")
    role = job_details.get("role", "Unknown Role")
    location = job_details.get("location", "UK")
    
    cv_path = r"C:\Users\tmalu\Downloads\Malundi_Christian_CV.pdf"
    cl_template_path = r"C:\Users\tmalu\Downloads\Malundi_Master_Cover_Letter_Template.docx"
    
    print(f"[*] AI Engine generating tailored documents for {role} at {company}...")
    
    # Extract CV text
    cv_text = extract_text_from_pdf(cv_path)
    if not cv_text:
        return "Error", "Master CV not found."

    # Create Prompt for LLM
    prompt = f"""
    You are an expert career coach. Tailor the following CV summary and generate a cover letter for a {role} position at {company} in {location}.
    
    Master CV Content:
    {cv_text[:2000]} # Truncated for token limits
    
    Job Context: {role} at {company} in {location}.
    
    Please provide:
    1. A tailored CV Summary (max 100 words).
    2. A full Professional Cover Letter.
    
    Format:
    ---SUMMARY---
    [Summary Text]
    ---LETTER---
    [Letter Text]
    """
    
    ai_output = ollama_generate(prompt)
    
    # Parse AI output
    try:
        summary_text = ai_output.split("---SUMMARY---")[1].split("---LETTER---")[0].strip()
        letter_text = ai_output.split("---LETTER---")[1].strip()
    except:
        summary_text = "AI generation error, please review manually."
        letter_text = ai_output

    output_dir = "Output"
    os.makedirs(output_dir, exist_ok=True)
    
    cv_out_path = os.path.join(output_dir, f"{company}_{role}_Tailored_CV.docx")
    cl_out_path = os.path.join(output_dir, f"{company}_{role}_CoverLetter.docx")
    
    # Save Tailored CV as Docx (easier to generate)
    doc_cv = Document()
    doc_cv.add_heading(f'Tailored CV for {role} at {company}', 0)
    doc_cv.add_heading('Professional Summary', level=1)
    doc_cv.add_paragraph(summary_text)
    doc_cv.add_heading('Note', level=1)
    doc_cv.add_paragraph("This summary was tailored by AI. Full CV details follow as per master document.")
    doc_cv.save(cv_out_path)
    
    # Save Cover Letter
    doc_cl = Document()
    doc_cl.add_heading('Cover Letter', 0)
    doc_cl.add_paragraph(letter_text)
    doc_cl.save(cl_out_path)
        
    return cv_out_path, cl_out_path

if __name__ == "__main__":
    # Test run
    print(generate_tailored_documents({
        "company": "TestCorp",
        "role": "Financial Analyst",
        "location": "London"
    }))
